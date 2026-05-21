"""
02_process_docs.py
==================
Reads raw ZIP files from S3 landing bucket, extracts and parses
3GPP spec documents (.docx / .pdf), performs structure-aware chunking,
generates per-chunk metadata via Claude Haiku, and writes chunk JSONs
to the processed S3 bucket.

Usage:
    python 02_process_docs.py               # process all unprocessed files
    python 02_process_docs.py --limit 5     # process first 5 (POC)
    python 02_process_docs.py --s3-key raw/38series/38300/38300-i40.zip
"""

import argparse
import io
import json
import os
import re
import time
import zipfile
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
import uuid

import boto3
import docx  # python-docx
import tiktoken

try:
    import fitz  # PyMuPDF for PDF parsing
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

from config import (
    AWS_REGION, S3_LANDING_BUCKET, S3_PROCESSED_BUCKET,
    DYNAMODB_TABLE, HAIKU_MODEL_ID,
    MAX_CHUNK_TOKENS, CHUNK_OVERLAP_TOKENS, MIN_CHUNK_TOKENS
)

s3     = boto3.client("s3",       region_name=AWS_REGION)
ddb    = boto3.resource("dynamodb", region_name=AWS_REGION)
table  = ddb.Table(DYNAMODB_TABLE)
bedrock = boto3.client("bedrock-runtime", region_name=AWS_REGION)
textract_client = boto3.client("textract", region_name=AWS_REGION)
enc    = tiktoken.get_encoding("cl100k_base")


# ─────────────────────────────────────────────────────────────────────────────
# Data structures
# ─────────────────────────────────────────────────────────────────────────────
@dataclass
class Section:
    number: str           # e.g. "5.3.2"
    title:  str           # e.g. "RRC State Machine"
    depth:  int           # heading level 1–6
    text:   str           # raw paragraph text
    tables: list[str]     = field(default_factory=list)  # markdown tables


@dataclass
class Chunk:
    chunk_id:     str
    doc_id:       str
    spec_number:  str     # e.g. "38300"
    spec_series:  str     # e.g. "38series"
    release:      str     # e.g. "Rel-18"
    section_path: str     # e.g. "5.3.2 RRC State Machine"
    doc_type:     str     # "TS" or "TR"
    chunk_text:   str
    token_count:  int
    source_s3_key: str
    # populated by metadata step
    summary:      str = ""
    keywords:     list = field(default_factory=list)
    hyp_questions: list = field(default_factory=list)


# ─────────────────────────────────────────────────────────────────────────────
# Parsing helpers
# ─────────────────────────────────────────────────────────────────────────────
CLAUSE_RE = re.compile(r"^(\d+(?:\.\d+){0,5})\s+(.+)")
RELEASE_RE = re.compile(r"[Rr]el(?:ease)?[-\s]?(\d+)", re.IGNORECASE)
TS_TR_RE   = re.compile(r"\b(TS|TR)\b")


def extract_docx_sections(doc: docx.Document) -> list[Section]:
    """Extract sections from a 3GPP .docx, preserving tables as markdown."""
    sections: list[Section] = []
    current: Optional[Section] = None

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            para = docx.text.paragraph.Paragraph(block, doc)
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            # Detect headings
            if "Heading" in style_name or re.match(r"Heading \d", style_name):
                depth = int(re.search(r"\d", style_name).group()) if re.search(r"\d", style_name) else 1
                m = CLAUSE_RE.match(text)
                number = m.group(1) if m else ""
                title  = m.group(2) if m else text
                if current:
                    sections.append(current)
                current = Section(number=number, title=title, depth=depth, text="")
            else:
                # Body paragraph
                if current is None:
                    current = Section(number="0", title="Introduction", depth=1, text="")
                current.text += f"\n{text}"

        elif tag == "tbl":
            tbl = docx.table.Table(block, doc)
            md = table_to_markdown(tbl)
            if current:
                current.tables.append(md)

    if current:
        sections.append(current)
    return sections


def table_to_markdown(tbl: docx.table.Table) -> str:
    """Convert a python-docx Table to a markdown string."""
    rows = []
    for i, row in enumerate(tbl.rows):
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def extract_pdf_sections(pdf_bytes: bytes) -> list[Section]:
    """Extract sections from a PDF using PyMuPDF."""
    if not HAS_PDF:
        return []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    sections: list[Section] = []
    current: Optional[Section] = None

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] != 0:  # text block only
                continue
            for line in block["lines"]:
                text = "".join(span["text"] for span in line["spans"]).strip()
                if not text:
                    continue
                # Detect headings by font size (>14pt) or bold
                max_size = max(span["size"] for span in line["spans"])
                is_bold = any("bold" in span.get("font", "").lower() for span in line["spans"])

                if max_size >= 14 or (is_bold and max_size >= 12):
                    m = CLAUSE_RE.match(text)
                    number = m.group(1) if m else ""
                    title = m.group(2) if m else text
                    if current:
                        sections.append(current)
                    current = Section(number=number, title=title, depth=1, text="")
                else:
                    if current is None:
                        current = Section(number="0", title="Introduction", depth=1, text="")
                    current.text += f"\n{text}"

    if current:
        sections.append(current)
    doc.close()
    return sections


def extract_pdf_with_textract(pdf_bytes: bytes, s3_key: str) -> list[Section]:
    """Fallback: use Amazon Textract for complex/scanned PDFs."""
    # Upload to S3 temporarily for Textract
    temp_key = f"temp/textract/{uuid.uuid4()}.pdf"
    s3.put_object(Bucket=S3_LANDING_BUCKET, Key=temp_key, Body=pdf_bytes)

    try:
        resp = textract_client.start_document_text_detection(
            DocumentLocation={"S3Object": {"Bucket": S3_LANDING_BUCKET, "Name": temp_key}}
        )
        job_id = resp["JobId"]

        # Wait for completion
        while True:
            result = textract_client.get_document_text_detection(JobId=job_id)
            status = result["JobStatus"]
            if status == "SUCCEEDED":
                break
            elif status == "FAILED":
                print(f"    ⚠ Textract failed for {s3_key}")
                return []
            time.sleep(2)

        # Collect all pages
        full_text = ""
        pages = [result]
        while "NextToken" in result:
            result = textract_client.get_document_text_detection(
                JobId=job_id, NextToken=result["NextToken"]
            )
            pages.append(result)

        for page in pages:
            for block in page["Blocks"]:
                if block["BlockType"] == "LINE":
                    full_text += block["Text"] + "\n"

        # Parse into sections
        sections: list[Section] = []
        current: Optional[Section] = None
        for line in full_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            m = CLAUSE_RE.match(line)
            if m and len(line) < 100:  # Short lines with clause numbers = headings
                if current:
                    sections.append(current)
                current = Section(number=m.group(1), title=m.group(2), depth=1, text="")
            else:
                if current is None:
                    current = Section(number="0", title="Document", depth=1, text="")
                current.text += f"\n{line}"

        if current:
            sections.append(current)
        return sections

    finally:
        # Cleanup temp file
        s3.delete_object(Bucket=S3_LANDING_BUCKET, Key=temp_key)


def parse_metadata_from_filename(filename: str) -> dict:
    """Extract spec number, release, doc type from filename like 38300-i40.docx"""
    stem = Path(filename).stem   # e.g. "38300-i40"
    parts = stem.split("-")
    spec_number = parts[0] if parts else stem

    # Release: 'i' prefix = Release-18, 'h' = 17, 'g' = 16, etc.
    release_map = {c: str(i) for i, c in enumerate("abcdefghijklmnopqrstuvwxyz", 10)}
    release = "unknown"
    if len(parts) > 1 and parts[1] and parts[1][0].lower() in release_map:
        release = f"Rel-{release_map[parts[1][0].lower()]}"

    doc_type = "TR" if spec_number.startswith(("21", "22", "25")) else "TS"
    series = spec_number[:2] + "series" if spec_number else ""
    return {"spec_number": spec_number, "release": release,
            "doc_type": doc_type, "series": series}


# ─────────────────────────────────────────────────────────────────────────────
# Structure-aware chunker
# ─────────────────────────────────────────────────────────────────────────────
def count_tokens(text: str) -> int:
    return len(enc.encode(text))


def build_breadcrumb(section: Section) -> str:
    """e.g. '[§5.3.2 RRC State Machine]'"""
    if section.number:
        return f"[§{section.number} {section.title}]"
    return f"[{section.title}]"


def chunk_section(section: Section, doc_meta: dict, doc_id: str,
                  source_s3_key: str) -> list[Chunk]:
    """Split one section into token-bounded chunks, never breaking a table."""
    chunks: list[Chunk] = []
    breadcrumb = build_breadcrumb(section)

    # Combine body text + tables (tables appended after paragraphs)
    full_text = section.text.strip()
    for t in section.tables:
        full_text += f"\n\n{t}"

    if not full_text.strip():
        return chunks

    # Split body text into sentences / paragraphs
    paragraphs = [p.strip() for p in full_text.split("\n") if p.strip()]

    current_paras: list[str] = []
    current_tokens = count_tokens(breadcrumb) + 4

    def flush(paras: list[str]) -> Optional[Chunk]:
        text = breadcrumb + "\n\n" + "\n".join(paras)
        tok = count_tokens(text)
        if tok < MIN_CHUNK_TOKENS:
            return None
        return Chunk(
            chunk_id     = str(uuid.uuid4()),
            doc_id       = doc_id,
            spec_number  = doc_meta["spec_number"],
            spec_series  = doc_meta["series"],
            release      = doc_meta["release"],
            section_path = f"{section.number} {section.title}".strip(),
            doc_type     = doc_meta["doc_type"],
            chunk_text   = text,
            token_count  = tok,
            source_s3_key = source_s3_key
        )

    for para in paragraphs:
        pt = count_tokens(para) + 1
        if current_tokens + pt > MAX_CHUNK_TOKENS and current_paras:
            c = flush(current_paras)
            if c:
                chunks.append(c)
            # Overlap: keep last few paragraphs
            overlap_paras = current_paras[-(CHUNK_OVERLAP_TOKENS // 20):]
            current_paras = overlap_paras
            current_tokens = count_tokens(breadcrumb) + sum(count_tokens(p) for p in overlap_paras) + 4
        current_paras.append(para)
        current_tokens += pt

    if current_paras:
        c = flush(current_paras)
        if c:
            chunks.append(c)

    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# Claude Haiku metadata generation
# ─────────────────────────────────────────────────────────────────────────────
def generate_chunk_metadata(chunk: Chunk) -> tuple[str, list[str], list[str]]:
    """Returns (summary, keywords, hypothetical_questions) via Claude Haiku."""
    prompt = f"""You are a 3GPP standards expert. Given the following chunk from a 3GPP specification, return ONLY valid JSON with exactly three fields:
- "summary": a 2-sentence technical summary
- "keywords": list of 5-8 technical terms (clause numbers, IE names, procedures)
- "hyp_questions": list of 3 questions a telecom engineer might ask that this chunk answers

Chunk:
{chunk.chunk_text[:1500]}

Return ONLY the JSON object, no markdown fences."""

    resp = bedrock.invoke_model(
        modelId=HAIKU_MODEL_ID,
        body=json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 512,
            "messages": [{"role": "user", "content": prompt}]
        }),
        contentType="application/json",
        accept="application/json"
    )
    raw = json.loads(resp["body"].read())["content"][0]["text"].strip()

    # Strip markdown fences if the model added them
    raw = re.sub(r"^```json\s*|\s*```$", "", raw.strip())
    try:
        data = json.loads(raw)
        return (
            data.get("summary", ""),
            data.get("keywords", []),
            data.get("hyp_questions", [])
        )
    except json.JSONDecodeError:
        return "", [], []


# ─────────────────────────────────────────────────────────────────────────────
# S3 + DynamoDB helpers
# ─────────────────────────────────────────────────────────────────────────────
def list_unprocessed_keys() -> list[str]:
    resp = table.scan(
        FilterExpression=boto3.dynamodb.conditions.Attr("processed").eq(False)
    )
    return [item["s3_key"] for item in resp.get("Items", [])]


def mark_processed(ftp_path_or_s3_key: str):
    # We stored file_path (ftp_path) as the PK; look up by s3_key
    resp = table.scan(
        FilterExpression=boto3.dynamodb.conditions.Attr("s3_key").eq(ftp_path_or_s3_key)
    )
    for item in resp.get("Items", []):
        table.update_item(
            Key={"file_path": item["file_path"]},
            UpdateExpression="SET #p = :t",
            ExpressionAttributeNames={"#p": "processed"},
            ExpressionAttributeValues={":t": True}
        )


def download_s3_zip(s3_key: str) -> io.BytesIO:
    obj = s3.get_object(Bucket=S3_LANDING_BUCKET, Key=s3_key)
    return io.BytesIO(obj["Body"].read())


def upload_chunks(chunks: list[Chunk], spec_number: str):
    key = f"chunks/{spec_number}/{chunks[0].doc_id}.jsonl"
    lines = "\n".join(json.dumps(asdict(c)) for c in chunks)
    s3.put_object(
        Bucket=S3_PROCESSED_BUCKET,
        Key=key,
        Body=lines.encode(),
        ContentType="application/jsonl"
    )
    return key


# ─────────────────────────────────────────────────────────────────────────────
# Main pipeline
# ─────────────────────────────────────────────────────────────────────────────
def process_s3_key(s3_key: str, skip_metadata: bool = False) -> int:
    """Download ZIP, parse docx inside, chunk, generate metadata, upload chunks."""
    print(f"  Processing: {s3_key}")
    zip_bytes = download_s3_zip(s3_key)
    filename  = Path(s3_key).name
    doc_meta  = parse_metadata_from_filename(filename)
    doc_id    = str(uuid.uuid4())
    all_chunks: list[Chunk] = []

    with zipfile.ZipFile(zip_bytes) as zf:
        docx_names = [n for n in zf.namelist() if n.lower().endswith(".docx")]
        pdf_names = [n for n in zf.namelist() if n.lower().endswith(".pdf")]

        if docx_names:
            inner_name = docx_names[0]
            with zf.open(inner_name) as f:
                doc = docx.Document(io.BytesIO(f.read()))
            sections = extract_docx_sections(doc)
            print(f"    Parsed {len(sections)} sections from {inner_name}")
        elif pdf_names and HAS_PDF:
            inner_name = pdf_names[0]
            with zf.open(inner_name) as f:
                pdf_bytes = f.read()
                sections = extract_pdf_sections(pdf_bytes)
            print(f"    Parsed {len(sections)} sections from {inner_name} (PDF)")
            # For short PDFs (presentations), merge all text into fewer sections
            if sections and all(count_tokens(s.text) < MIN_CHUNK_TOKENS for s in sections):
                merged_text = "\n".join(f"[{s.title}]\n{s.text}" for s in sections if s.text.strip())
                if count_tokens(merged_text) < MIN_CHUNK_TOKENS:
                    # PyMuPDF got too little text — fallback to Textract
                    print(f"    ⚠ PyMuPDF insufficient, falling back to Textract...")
                    sections = extract_pdf_with_textract(pdf_bytes, s3_key)
                    print(f"    Textract extracted {len(sections)} sections")
                else:
                    sections = [Section(number="0", title=sections[0].title, depth=1, text=merged_text)]
        else:
            print(f"    ⚠ No .docx or .pdf found in {filename}, skipping")
            return 0

        for section in sections:
            chunks = chunk_section(section, doc_meta, doc_id, s3_key)
            all_chunks.extend(chunks)

    print(f"    Created {len(all_chunks)} chunks" + (" — generating metadata …" if not skip_metadata else " — skipping metadata"))
    if not skip_metadata:
        for i, chunk in enumerate(all_chunks):
            try:
                summary, kws, hqs = generate_chunk_metadata(chunk)
                chunk.summary       = summary
                chunk.keywords      = kws
                chunk.hyp_questions = hqs
            except Exception as e:
                print(f"      Metadata error on chunk {i}: {e}")
            time.sleep(0.3)   # Haiku rate limit courtesy

    if not all_chunks:
        print(f"    ⚠ No chunks created (text too short), skipping")
        mark_processed(s3_key)
        return 0

    chunks_key = upload_chunks(all_chunks, doc_meta["spec_number"])
    mark_processed(s3_key)
    print(f"    ✓ {len(all_chunks)} chunks → s3://{S3_PROCESSED_BUCKET}/{chunks_key}")
    return len(all_chunks)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit",   type=int, default=None)
    parser.add_argument("--s3-key",  type=str, default=None)
    parser.add_argument("--skip-metadata", action="store_true",
                        help="Skip Haiku metadata generation (much faster)")
    args = parser.parse_args()

    if args.s3_key:
        keys = [args.s3_key]
    else:
        keys = list_unprocessed_keys()
        print(f"Found {len(keys)} unprocessed files in DynamoDB manifest")
        if args.limit:
            keys = keys[:args.limit]

    total_chunks = 0
    for s3_key in keys:
        try:
            n = process_s3_key(s3_key, skip_metadata=args.skip_metadata)
            total_chunks += n
        except Exception as e:
            print(f"  ✗ Failed {s3_key}: {e}")

    print(f"\n── Done ──────────────────────────────────────────────────────")
    print(f"  Processed files : {len(keys)}")
    print(f"  Total chunks    : {total_chunks}")
    print(f"\nNext step: python 03_embed_and_index.py")


if __name__ == "__main__":
    main()
